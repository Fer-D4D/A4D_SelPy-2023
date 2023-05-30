import time
import calendar
from datetime import datetime

import docx
from docx.shared import Inches, Pt
from openpyxl.reader.excel import load_workbook
from selenium import webdriver
from selenium.common import NoSuchElementException, InvalidSelectorException, TimeoutException, \
    ElementNotVisibleException, ElementNotSelectableException
from selenium.webdriver import ActionChains
from selenium.webdriver.edge.service import Service as EdgeService
from selenium.webdriver.support.select import Select
from selenium.webdriver.support.wait import WebDriverWait
from webdriver_manager.microsoft import EdgeChromiumDriverManager
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.firefox.service import Service as FirefoxService
from webdriver_manager.firefox import GeckoDriverManager
from selenium.webdriver.chrome.service import Service as BraveService
from webdriver_manager.chrome import ChromeDriverManager
from webdriver_manager.core.utils import ChromeType
from selenium.webdriver.support import expected_conditions as ec

from Team.Fer.core.utils import Utils, timer


def get_by_string(locator_definition):
    if "CSS:" in locator_definition:
        return ["css selector", locator_definition.replace("CSS:", "")]
    if "ID:" in locator_definition:
        return ["css selector", locator_definition.replace("ID:", "#")]
    if "XPATH:" in locator_definition:
        return ["xpath", locator_definition.replace("XPATH:", "")]
    if "NAME:" in locator_definition:
        return ["xpath", locator_definition.replace("NAME:", "//*[@name='") + "']"]
    if "LINK_TEXT:" in locator_definition:
        return ["xpath", locator_definition.replace("LINK_TEXT:", "//a[contains(text(),'") + "']"]


def map_to_boolean(word):
    if "on" in word.lower():
        return True
    else:
        return False


def check_for_none_type(element):
    if element is None:
        return False
    else:
        return True


def get_time_stamp():
    return str(calendar.timegm(time.gmtime()))


def boolean_mapper(boolean_input, value_for_true, value_for_false):
    if boolean_input:
        return value_for_true
    return value_for_false


def get_current_date():
    now = datetime.now()
    return f"{now.month:02d}/{now.day:02d}/{now.year} {now.hour:02d}:{now.minute:02d}"


def get_last_row(worksheet):
    for row_num in range(1, worksheet.max_row):
        if worksheet.cell(row=row_num, column=1).value is None:
            return row_num


class TinyCore(Utils):
    DRIVER = None
    BROWSER = 'chrome'
    PAGE_TIME_OUT = 15
    FLUENT_WAIT_TIMEOUT = 5
    FLUENT_WAIT_FREQ = 1
    DEFAULT_TIME_TO_WASTE = 1

    VIEWER_MODE = False
    VIEWER_MODE_TIME = 1
    VERBOSE_MODE = False
    HIGHLIGHT_MODE = False

    TEST_STEP_COUNTER = 0

    FLOW_CONTROL_FLAG = False
    SS_FAIL_MODE = False
    DOC_OBJECT = None
    XLSX_OBJ = None
    TEST_RUN = 0

    TEST_DOCUMENTS_PATH = 'C:/Automation/docx'
    TEST_DATA_PATH = 'C:/Automation/TestData'
    DEFAULT_SCREEN_SHOT_PATH = 'C:/SS_Automation'
    DEFAULT_TRUSTED_KEY_SELECTOR = "XPATH://body"
    DEFAULT_TEST_SS_NAME = "SS_test"
    GENERIC_FAIL_MESSAGE = "Test step failed, please check."

    HIGHLIGHT_COLOR = "red"
    HIGHLIGHT_BORDER = 4
    HIGHLIGHT_DURATION = 1

    def __init__(self, browser='chrome', viewer_mode="Viewer-Mode-OFF", verbose_mode="Verbose-Mode-OFF",
                 highlight_mode="Highlight-Mode-OFF"):
        self.BROWSER = browser.lower()
        self.VIEWER_MODE = map_to_boolean(viewer_mode)
        self.VERBOSE_MODE = map_to_boolean(verbose_mode)
        self.HIGHLIGHT_MODE = map_to_boolean(highlight_mode)

    def define_webdriver(self):
        if self.BROWSER == "edge":
            return webdriver.Edge(service=EdgeService(EdgeChromiumDriverManager().install()))
        if self.BROWSER == "firefox":
            return webdriver.Firefox(service=FirefoxService(GeckoDriverManager().install()))
        if self.BROWSER == "brave":
            return webdriver.Chrome(service=BraveService(ChromeDriverManager(chrome_type=ChromeType.BRAVE).install()))
        return webdriver.Chrome(service=ChromeService(ChromeDriverManager().install()))

    def start_driver(self):
        self.DRIVER = self.define_webdriver()

    def get_driver(self):
        return self.DRIVER

    @timer
    def launch_site(self, base_url, trusted_key_selector=DEFAULT_TRUSTED_KEY_SELECTOR, dynamic_selector_text=""):
        self.DRIVER = self.define_webdriver()
        self.DRIVER.get(base_url)
        self.DRIVER.maximize_window()
        self.is_page_ready(trusted_key_selector, dynamic_selector_text)

    def get_element(self, selector_definition, dynamic_selector_text="", waiting_time=FLUENT_WAIT_TIMEOUT):
        wait = WebDriverWait(self.DRIVER, waiting_time, poll_frequency=self.FLUENT_WAIT_FREQ,
                             ignored_exceptions=[ElementNotVisibleException, ElementNotSelectableException])
        by_string = self.create_required_bystring(selector_definition, dynamic_selector_text)
        try:
            element = wait.until(ec.element_to_be_clickable((by_string["By"], by_string["custom_selector"])))
            self.highlight_mode(element)
            self.verbose_mode("<" + selector_definition + "> Selector found!")
            return element
        except NoSuchElementException:
            self.verbose_mode("<" + selector_definition + "> Selector not found, please check it out")
            return None
        except InvalidSelectorException:
            self.verbose_mode("<" + selector_definition + "> Selector is invalid please check it out.")
            return None
        except TimeoutException:
            self.verbose_mode("<" + selector_definition + "> Selector not found, please check it out")
            return None

    def is_page_ready(self, trusted_key_selector=DEFAULT_TRUSTED_KEY_SELECTOR, dynamic_selector_text=""):
        print(self.validate_object(self.DRIVER))
        print(self.create_required_bystring(trusted_key_selector, dynamic_selector_text,
                                            self.PAGE_TIME_OUT))
        if self.validate_object(self.DRIVER) and self.validate_object(
                self.get_element(self.create_required_bystring(trusted_key_selector), dynamic_selector_text,
                                 self.PAGE_TIME_OUT)):
            self.verbose_mode("The provided trusted key element <" + trusted_key_selector + "> was found, so we can "
                                                                                            "assume that the page has "
                                                                                            "been "
                                                                                            "loaded.")
            return True
        else:
            self.verbose_mode("Timed out waiting for page to load, please check the provided trusted key element <"
                              + trusted_key_selector + ">")
            return False

    def viewer_mode(self):
        if self.VIEWER_MODE:
            self.waste_some_time(self.VIEWER_MODE_TIME)

    def set_viewer_mode(self, viewer_mode="Viewer-Mode-OFF"):
        self.VIEWER_MODE = map_to_boolean(viewer_mode)

    def set_browser(self, browser='chrome'):
        self.BROWSER = browser.lower()

    def set_verbose_mode(self, verbose_mode="Verbose-Mode-OFF"):
        self.VERBOSE_MODE = map_to_boolean(verbose_mode)

    def set_highlight_mode(self, highlight_mode="Verbose-Mode-OFF"):
        self.HIGHLIGHT_MODE = map_to_boolean(highlight_mode)

    def set_driver(self, driver):
        self.DRIVER = driver

    def update_flow_control_flag(self, boolean_state):
        self.FLOW_CONTROL_FLAG = boolean_state

    def set_ss_fail_mode(self, ss_fail_mode="SS-Fail-Mode-OFF"):
        self.SS_FAIL_MODE = map_to_boolean(ss_fail_mode)

    def reset_step_counter(self):
        self.TEST_STEP_COUNTER = 0

    def reset_test_run(self):
        self.TEST_RUN = 0

    def switch_flow_control_flag(self):
        if self.get_flow_control_flag_status():
            self.update_flow_control_flag(False)
        else:
            self.update_flow_control_flag(True)

    def get_flow_control_flag_status(self):
        return self.FLOW_CONTROL_FLAG

    def update_test_run(self):
        self.TEST_RUN = self.TEST_RUN + 1
        return self.TEST_RUN

    def verbose_mode(self, verbose_msg):
        if self.VERBOSE_MODE:
            print(verbose_msg)

    def highlight_mode(self, element):
        if self.HIGHLIGHT_MODE:
            self.highlight_element(element)

    def highlight_element(self, element, restore_style=True):
        original_style = element.get_attribute('style')
        self.apply_style(element, "border: {0}px solid {1};".format(self.HIGHLIGHT_BORDER, self.HIGHLIGHT_COLOR))
        if restore_style:
            time.sleep(self.HIGHLIGHT_DURATION)
            self.apply_style(element, original_style)

    def apply_style(self, element, style):
        self.DRIVER.execute_script("arguments[0].setAttribute('style', arguments[1]);", element, style)

    def go_to_element(self, web_element):
        if check_for_none_type(web_element):
            ActionChains(self.DRIVER).move_to_element(web_element).perform()

    def do_click(self, locator_definition):
        element = self.get_element(locator_definition)
        if check_for_none_type(element):
            element.click()
            self.viewer_mode()
            self.verbose_mode("Element <" + locator_definition + "> successfully clicked!")
            return True
        else:
            self.verbose_mode("Element <" + locator_definition + "> not clicked")
            return False

    def fill_input_text(self, locator_definition, text):
        element = self.get_element(locator_definition)
        if check_for_none_type(element):
            element.click()
            element.clear()
            element.send_keys(text)
            self.viewer_mode()
            self.verbose_mode("Element <" + locator_definition + "> successfully filled in!")
            return True
        else:
            self.verbose_mode("Element <" + locator_definition + "> not filled in")
            return False

    def get_element_inner_text(self, locator_definition):
        element = self.get_element(locator_definition)
        if check_for_none_type(element):
            inner_text = element.text
            self.viewer_mode()
            self.verbose_mode("The element <" + locator_definition + "> retrieves this text ->  " + inner_text)
            return inner_text
        else:
            self.verbose_mode("The element <" + locator_definition + "> retrieves no text.")
            return None

    def select_value_from_dropdown(self, locator_definition, desired_value):
        element = self.get_element(locator_definition)
        if check_for_none_type(element):
            Select(element).select_by_visible_text(desired_value)
            self.viewer_mode()
            self.verbose_mode("Desired value selected from Element <<" + locator_definition + ">")
            return True
        else:
            self.verbose_mode("Unable to select desire value from Element <" + locator_definition + ">")
            return False

    def page_back(self):
        self.viewer_mode()
        self.DRIVER.execute_script("window.history.go(-1)")

    def switch_browser_tab(self):
        windows = self.DRIVER.window_handles
        self.DRIVER.switch_to.window(windows[1])
        self.viewer_mode()
        self.DRIVER.close()
        self.DRIVER.switch_to.window(windows[0])

    def get_extended_screenshot(self, screen_shot_title, main_content=DEFAULT_TRUSTED_KEY_SELECTOR):
        # Where to save the picture
        # Video title
        gmt = time.gmtime()
        ts = calendar.timegm(gmt)

        try:
            # We try to get the top-level component in which all out page is its children
            # This is different for each website
            elem = self.get_element(main_content)
            # Get the height of the element, and adding some height just to be sage
            total_height = elem.size['height'] + 1000
            # Set the window size - what is the size of our screenshot
            # The width is hardcoded because the screensize is fixed for each computer
            self.DRIVER.set_window_size(1920, total_height)
            # Wait for 2 seconds
            self.waste_some_time(2)
            # Take the screenshot
            elem.screenshot(f'{self.DEFAULT_SCREEN_SHOT_PATH}/{screen_shot_title}_{ts}.png')
            return f'{screen_shot_title}_{ts}.png'
        except SystemError:
            print('Take screenshot error at' + screen_shot_title)

    def get_full_screenshot(self, screen_shot_title=DEFAULT_TEST_SS_NAME, main_content=DEFAULT_TRUSTED_KEY_SELECTOR):
        try:
            ts = get_time_stamp()
            elem = self.get_element(main_content)
            elem.screenshot(f'{self.DEFAULT_SCREEN_SHOT_PATH}/{screen_shot_title}_{ts}.png')
            return f'{screen_shot_title}_{ts}.png'
        except SystemError:
            print('Take screenshot error at' + screen_shot_title)

    def get_emphasis_screenshot(self, screen_shot_title, locator_definition):
        gmt = time.gmtime()
        ts = calendar.timegm(gmt)
        try:
            elem = self.get_element(self.DEFAULT_TRUSTED_KEY_SELECTOR)

            # Get the height of the element, and adding some height just to be sage
            total_height = elem.size['height'] + 1000
            # Set the window size - what is the size of our screenshot
            # The width is hardcoded because the screensize is fixed for each computer
            self.DRIVER.set_window_size(1920, total_height)
            # Wait for 2 seconds
            self.highlight_element(self.get_element(locator_definition), False)
            # waste_some_time(2)

            # Take the screenshot
            elem.screenshot(f'{self.DEFAULT_SCREEN_SHOT_PATH}/{screen_shot_title}_{ts}.png')
        except SystemError:
            print('Take screenshot error at' + screen_shot_title)

    def get_basic_screenshot(self, screen_shot_title=DEFAULT_TEST_SS_NAME):
        ts = get_time_stamp()
        name = f'{screen_shot_title}_{ts}.png'
        self.DRIVER.save_screenshot(f'{self.DEFAULT_SCREEN_SHOT_PATH}/{name}')
        return name

    def get_elements_list(self, locator_definition):
        by_string = get_by_string(locator_definition)
        return self.DRIVER.find_elements(by_string[0], by_string[1])

    def get_number_of_elements(self, locator_definition):
        return len(self.get_elements_list(locator_definition))

    def compare_element_inner_text(self, locator_definition, expected_text):
        return self.get_element_inner_text(locator_definition) == expected_text

    def create_test_doc(self):
        self.reset_step_counter()
        return docx.Document()

    def add_step_to_test_doc(self, doc_obj, step_definition, associated_screenshot=None, last_step=False):
        self.TEST_STEP_COUNTER = self.TEST_STEP_COUNTER + 1
        doc_obj.add_heading(f"Step {self.TEST_STEP_COUNTER:02d}: {step_definition}")
        if associated_screenshot is not None:
            doc_obj.add_picture(f"{self.DEFAULT_SCREEN_SHOT_PATH}/{associated_screenshot}", width=Inches(6))
        if not last_step and self.TEST_STEP_COUNTER % 2 == 0:
            doc_obj.add_page_break()

    def save_test_doc(self, doc_obj, doc_name):
        doc_obj.save(f"{self.TEST_DOCUMENTS_PATH}/{doc_name}_{get_time_stamp()}.docx")

    @staticmethod
    def compare_lists(list_x, list_y):
        return sorted(list_x) == sorted(list_y)

    def safe_to_proceed(self, key_locator_definition):
        return self.get_number_of_elements(key_locator_definition) > 0

    def start_test_doc(self):
        self.reset_step_counter()
        self.DOC_OBJECT = docx.Document()

    def add_cover_page_old(self, summary, env_details, tested_by):
        style = self.DOC_OBJECT.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(18)

        p = self.DOC_OBJECT.add_paragraph()
        p.style = self.DOC_OBJECT.styles['Normal']
        p.add_run("SUMMARY: ").bold = True
        p.add_run(summary).italic = True
        p = self.DOC_OBJECT.add_paragraph()
        p.add_run("Environment: ").bold = True
        p.add_run(env_details).italic = True
        p = self.DOC_OBJECT.add_paragraph()
        p.add_run("Tested By: ").bold = True
        p.add_run(tested_by).italic = True
        self.DOC_OBJECT.add_page_break()

    def add_cover_page(self, cover_item_list):
        self.add_doc_style("Normal", "Calibri", 22)
        for item, definition in cover_item_list.items():
            self.add_paragraph_docx("Normal", item, definition.replace("URT", get_current_date()), ": ")
        self.DOC_OBJECT.add_page_break()

    def add_doc_style(self, style_name, font_family, font_size):
        font = self.DOC_OBJECT.styles[style_name].font
        font.name = font_family
        font.size = Pt(font_size)

    def add_paragraph_docx(self, style, left_text="", right_text="", separator=""):
        p = self.DOC_OBJECT.add_paragraph()
        p.style = self.DOC_OBJECT.styles[style]
        p.add_run(left_text + separator).bold = True
        p.add_run(right_text).italic = True

    def document_assert_results(self, boolean_result, assert_message, assert_failed_message=GENERIC_FAIL_MESSAGE,
                                last_assert=False):
        self.update_flow_control_flag(boolean_result)
        self.TEST_STEP_COUNTER = self.TEST_STEP_COUNTER + 1
        self.DOC_OBJECT.add_heading(f"Step {self.TEST_STEP_COUNTER:02d} - {assert_message}: "
                                    f"{boolean_mapper(boolean_result, 'PASSED', 'FAILED')}")
        if not boolean_result:
            self.add_doc_style("Quote", "Calibri", 12)
            self.add_paragraph_docx("Quote", "", f"Failure reason: {assert_failed_message}")
        if not self.SS_FAIL_MODE:
            self.DOC_OBJECT.add_picture(f"{self.DEFAULT_SCREEN_SHOT_PATH}/"
                                        f"{self.get_basic_screenshot()}", width=Inches(6))
        if self.SS_FAIL_MODE and last_assert and self.TEST_STEP_COUNTER % 2 == 0:
            self.DOC_OBJECT.add_page_break()

    def save_doc_results(self, doc_name):
        self.DOC_OBJECT.save(f"{self.TEST_DOCUMENTS_PATH}/{doc_name}_{get_time_stamp()}.docx")

    # openpixl fun

    def open_xlsx(self, xlsx_file_name):
        self.XLSX_OBJ = load_workbook(filename=f"{self.TEST_DATA_PATH}/{xlsx_file_name}")

    def get_test_data_from_xlsx(self):
        sheet = self.XLSX_OBJ.active
        first_row = True
        test_data_rows = []
        headers = None
        for row in sheet.iter_rows(min_row=1,
                                   max_row=get_last_row(sheet),
                                   max_col=8,
                                   values_only=True):
            # print(row)
            if first_row:
                first_row = False
                headers = row
            else:
                test_dic = {}
                for iterator in range(0, len(headers)):
                    test_dic.update({headers[iterator]: row[iterator]})
                test_data_rows.append(test_dic)
        return test_data_rows

    def is_checkbox_checked(self, locator_definition):
        return self.get_element(locator_definition).is_selected()

    def check_checkbox(self, locator_definition):
        if not self.is_checkbox_checked(locator_definition):
            self.do_click(locator_definition)

    def uncheck_checkbox(self, locator_definition):
        if self.is_checkbox_checked(locator_definition):
            self.do_click(locator_definition)

    def compare_string_against_selector_text(self, locator_definition, string_to_compare):
        return self.get_element_inner_text(locator_definition) == string_to_compare

    @staticmethod
    def compare_string_against_web_element_text(web_element, string_to_compare):
        return web_element.text == string_to_compare

    def token_replace(self, locator_definition, new_value):
        return locator_definition.replace(self.GENERIC_TOKEN, new_value)
