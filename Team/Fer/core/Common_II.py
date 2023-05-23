import time
import calendar;
from functools import wraps

import docx
from Screenshot import Screenshot
from docx.shared import Inches, Pt
from selenium import webdriver
from selenium.common import NoSuchElementException, InvalidSelectorException, TimeoutException, \
    ElementNotVisibleException, ElementNotSelectableException
from selenium.webdriver import ActionChains
from selenium.webdriver.edge.service import Service as EdgeService
from selenium.webdriver.support.select import Select
from selenium.webdriver.support.wait import WebDriverWait
from webdriver_manager.microsoft import EdgeChromiumDriverManager
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.firefox.service import Service as FirefoxService
from webdriver_manager.firefox import GeckoDriverManager
from selenium.webdriver.chrome.service import Service as BraveService
from webdriver_manager.chrome import ChromeDriverManager
from webdriver_manager.core.utils import ChromeType
from selenium.webdriver.support import expected_conditions as ec


def waste_some_time(waiting_time=1):
    time.sleep(waiting_time)


def timer(f):
    def wrapper(*args, **kwargs):
        start_time = time.time()
        result = f(*args, **kwargs)
        stop_time = time.time()
        dt = stop_time - start_time
        print(f"Δt = {dt}")

    return wrapper


def get_by_string(locator_definition):
    if "CSS:" in locator_definition:
        return ["css selector", locator_definition.replace("CSS:", "")]
    if "ID:" in locator_definition:
        return ["css selector", locator_definition.replace("ID:", "#")]
    if "XPATH:" in locator_definition:
        return ["xpath", locator_definition.replace("XPATH:", "")]
    if "NAME:" in locator_definition:
        return ["xpath", locator_definition.replace("NAME:", "//*[@name='") + "']"]
    if "LINK_TEXT:" in locator_definition:
        return ["xpath", locator_definition.replace("LINK_TEXT:", "//a[contains(text(),'") + "']"]


def map_to_boolean(word):
    if "on" in word.lower():
        return True
    else:
        return False


def check_for_none_type(element):
    if element is None:
        return False
    else:
        return True


def get_time_stamp():
    return str(calendar.timegm(time.gmtime()))


def boolean_mapper(boolean_input, value_for_true, value_for_false):
    if boolean_input:
        return value_for_true
    return value_for_false


class TinyCore:
    DRIVER = None
    BROWSER = 'chrome'
    PAGE_TIME_OUT = 15
    FLUENT_WAIT_TIMEOUT = 3
    FLUENT_WAIT_FREQ = 1
    DEFAULT_TIME_TO_WASTE = 1
    HIGHLIGHT_COLOR = "red"
    HIGHLIGHT_BORDER = 4
    HIGHLIGHT_DURATION = 1
    VIEWER_MODE = False
    VIEWER_MODE_TIME = 1
    VERBOSE_MODE = False
    HIGHLIGHT_MODE = False
    DEFAULT_TRUSTED_KEY_ELEMENT = "XPATH://body"
    SELENIUM_WEB_ELEMENT_TYPE = 'selenium.webdriver.remote.webelement.WebElement'
    DEFAULT_SCREEN_SHOT_PATH = 'C:/SS_Automation'
    TEST_STEP_COUNTER = 0
    DEFAULT_TEST_SS_NAME = "SS_test"
    FLOW_CONTROL_FLAG = False
    SS_FAIL_MODE = False
    DOC_OBJECT = None
    TEST_RUN = 0

    TEST_DOCUMENTS_PATH = 'C:/Automation/docx'

    def __init__(self, browser='chrome', viewer_mode="Viewer-Mode-OFF", verbose_mode="Verbose-Mode-OFF",
                 highlight_mode="Highlight-Mode-OFF"):
        self.BROWSER = browser.lower()
        self.VIEWER_MODE = map_to_boolean(viewer_mode)
        self.VERBOSE_MODE = map_to_boolean(verbose_mode)
        self.HIGHLIGHT_MODE = map_to_boolean(highlight_mode)
        self.FLOW_CONTROL_FLAG = False

    def viewer_mode(self):
        if self.VIEWER_MODE:
            waste_some_time(self.VIEWER_MODE_TIME)

    def set_viewer_mode(self, viewer_mode="Viewer-Mode-OFF"):
        self.VIEWER_MODE = map_to_boolean(viewer_mode)

    def set_browser(self, browser='chrome'):
        self.BROWSER = browser.lower()

    def set_verbose_mode(self, verbose_mode="Verbose-Mode-OFF"):
        self.VERBOSE_MODE = map_to_boolean(verbose_mode)

    def set_highlight_mode(self, highlight_mode="Verbose-Mode-OFF"):
        self.HIGHLIGHT_MODE = map_to_boolean(highlight_mode)

    def set_driver(self, driver):
        self.DRIVER = driver

    def get_driver(self):
        return self.DRIVER

    def update_flow_control_flag(self, boolean_state):
        self.FLOW_CONTROL_FLAG = boolean_state

    def set_ss_fail_mode(self, ss_fail_mode="SS-Fail-Mode-OFF"):
        self.SS_FAIL_MODE = map_to_boolean(ss_fail_mode)

    def reset_step_counter(self):
        self.TEST_STEP_COUNTER = 0

    def reset_test_run(self):
        self.TEST_RUN = 0

    def update_test_run(self):
        self.TEST_RUN = self.TEST_RUN + 1
        return self.TEST_RUN

    def verbose_mode(self, verbose_msg):
        if self.VERBOSE_MODE:
            print(verbose_msg)

    def highlight_mode(self, element):
        if self.HIGHLIGHT_MODE:
            self.highlight_element(element)

    def highlight_element(self, element, restore_style=True):
        original_style = element.get_attribute('style')
        self.apply_style(element, "border: {0}px solid {1};".format(self.HIGHLIGHT_BORDER, self.HIGHLIGHT_COLOR))
        if restore_style:
            time.sleep(self.HIGHLIGHT_DURATION)
            self.apply_style(element, original_style)

    def apply_style(self, element, style):
        self.DRIVER.execute_script("arguments[0].setAttribute('style', arguments[1]);", element, style)

    def get_element(self, locator_definition):
        # self.viewer_mode()
        wait = WebDriverWait(self.DRIVER, self.FLUENT_WAIT_TIMEOUT, poll_frequency=self.FLUENT_WAIT_FREQ,
                             ignored_exceptions=[ElementNotVisibleException, ElementNotSelectableException])
        by_string = get_by_string(locator_definition)
        try:
            # element = self.driver.find_element(by_string[0], by_string[1])
            element = wait.until(ec.element_to_be_clickable((by_string[0], by_string[1])))
            self.highlight_mode(element)
            self.verbose_mode("<" + locator_definition + "> Selector found!")
            return element
        except NoSuchElementException:
            self.verbose_mode("<" + locator_definition + "> Selector not found, please check it out")
            return None
        except InvalidSelectorException:
            self.verbose_mode("<" + locator_definition + "> Selector is invalid please check it out.")
            return None
        except TimeoutException:
            self.verbose_mode("<" + locator_definition + "> Selector not found, please check it out")
            return None

    def get_webdriver(self):
        if self.BROWSER == "edge":
            return webdriver.Edge(service=EdgeService(EdgeChromiumDriverManager().install()))
        if self.BROWSER == "firefox":
            return webdriver.Firefox(service=FirefoxService(GeckoDriverManager().install()))
        if self.BROWSER == "brave":
            return webdriver.Chrome(service=BraveService(ChromeDriverManager(chrome_type=ChromeType.BRAVE).install()))
        return webdriver.Chrome(service=ChromeService(ChromeDriverManager().install()))

    def wait_for_page_safe_load(self, trusted_key_element=DEFAULT_TRUSTED_KEY_ELEMENT):
        by_string = get_by_string(trusted_key_element)
        try:
            WebDriverWait(self.DRIVER, self.PAGE_TIME_OUT).until(ec.presence_of_element_located((by_string[0],
                                                                                                 by_string[1])))
            self.verbose_mode("The provided trusted key element <" + trusted_key_element + "> was found, so we can "
                                                                                           "assume that the page has been "
                                                                                           "loaded.")
            return True
        except TimeoutException:
            self.verbose_mode("Timed out waiting for page to load, please check the provided trusted key element <"
                              + trusted_key_element + ">")
            return False

    def launch_site(self, base_url, anchor_locator_definition=DEFAULT_TRUSTED_KEY_ELEMENT):
        self.DRIVER = self.get_webdriver()
        self.DRIVER.get(base_url)
        self.DRIVER.maximize_window()
        self.wait_for_page_safe_load(anchor_locator_definition)
        return self.DRIVER

    def go_to_element(self, web_element):
        if check_for_none_type(web_element):
            ActionChains(self.DRIVER).move_to_element(web_element).perform()

    def do_click(self, locator_definition):
        element = self.get_element(locator_definition)
        if check_for_none_type(element):
            element.click()
            self.viewer_mode()
            self.verbose_mode("Element <" + locator_definition + "> successfully clicked!")
            return True
        else:
            self.verbose_mode("Element <" + locator_definition + "> not clicked")
            return False

    def fill_input_text(self, locator_definition, text):
        element = self.get_element(locator_definition)
        if check_for_none_type(element):
            element.click()
            element.clear()
            element.send_keys(text)
            self.viewer_mode()
            self.verbose_mode("Element <" + locator_definition + "> successfully filled in!")
            return True
        else:
            self.verbose_mode("Element <" + locator_definition + "> not filled in")
            return False

    def get_element_inner_text(self, locator_definition):
        element = self.get_element(locator_definition)
        if check_for_none_type(element):
            inner_text = element.text
            self.viewer_mode()
            self.verbose_mode("The element <" + locator_definition + "> retrieves this text ->  " + inner_text)
            return inner_text
        else:
            self.verbose_mode("The element <" + locator_definition + "> retrieves no text.")
            return None

    def select_value_from_dropdown(self, locator_definition, desired_value):
        element = self.get_element(locator_definition)
        if check_for_none_type(element):
            Select(element).select_by_visible_text(desired_value)
            self.viewer_mode()
            self.verbose_mode("Desired value selected from Element <<" + locator_definition + ">")
            return True
        else:
            self.verbose_mode("Unable to select desire value from Element <" + locator_definition + ">")
            return False

    def page_back(self):
        self.viewer_mode()
        self.DRIVER.execute_script("window.history.go(-1)")

    def switch_browser_tab(self):
        windows = self.DRIVER.window_handles
        self.DRIVER.switch_to.window(windows[1])
        self.viewer_mode()
        self.DRIVER.close()
        self.DRIVER.switch_to.window(windows[0])

    def get_extended_screenshot(self, screen_shot_title, main_content=DEFAULT_TRUSTED_KEY_ELEMENT):
        # Where to save the picture
        # Video title
        gmt = time.gmtime()
        ts = calendar.timegm(gmt)

        try:
            # We try to get the top-level component in which all out page is its children
            # This is different for each website
            elem = self.get_element(main_content)
            # Get the height of the element, and adding some height just to be sage
            total_height = elem.size['height'] + 1000
            # Set the window size - what is the size of our screenshot
            # The width is hardcoded because the screensize is fixed for each computer
            self.DRIVER.set_window_size(1920, total_height)
            # Wait for 2 seconds
            waste_some_time(2)
            # Take the screenshot
            elem.screenshot(f'{self.DEFAULT_SCREEN_SHOT_PATH}/{screen_shot_title}_{ts}.png')
            return f'{screen_shot_title}_{ts}.png'
        except SystemError as err:
            print('Take screenshot error at' + screen_shot_title)

    def get_full_screenshot(self, screen_shot_title=DEFAULT_TEST_SS_NAME, main_content=DEFAULT_TRUSTED_KEY_ELEMENT):
        try:
            ts = get_time_stamp()
            elem = self.get_element(main_content)
            elem.screenshot(f'{self.DEFAULT_SCREEN_SHOT_PATH}/{screen_shot_title}_{ts}.png')
            return f'{screen_shot_title}_{ts}.png'
        except SystemError as err:
            print('Take screenshot error at' + screen_shot_title)

    def get_emphasis_screenshot(self, screen_shot_title, locator_definition):
        gmt = time.gmtime()
        ts = calendar.timegm(gmt)
        try:
            elem = self.get_element(self.DEFAULT_TRUSTED_KEY_ELEMENT)

            # Get the height of the element, and adding some height just to be sage
            total_height = elem.size['height'] + 1000
            # Set the window size - what is the size of our screenshot
            # The width is hardcoded because the screensize is fixed for each computer
            self.DRIVER.set_window_size(1920, total_height)
            # Wait for 2 seconds
            self.highlight_element(self.get_element(locator_definition), False)
            # waste_some_time(2)

            # Take the screenshot
            elem.screenshot(f'{self.DEFAULT_SCREEN_SHOT_PATH}/{screen_shot_title}_{ts}.png')
        except SystemError as err:
            print('Take screenshot error at' + screen_shot_title)

    def get_basic_screenshot(self, screen_shot_title=DEFAULT_TEST_SS_NAME):
        ts = get_time_stamp()
        name = f'{screen_shot_title}_{ts}.png'
        self.DRIVER.save_screenshot(f'{self.DEFAULT_SCREEN_SHOT_PATH}/{name}')
        return name

    def get_elements_list(self, locator_definition):
        by_string = get_by_string(locator_definition)
        return self.DRIVER.find_elements(by_string[0], by_string[1])

    def get_number_of_elements(self, locator_definition):
        return len(self.get_elements_list(locator_definition))

    def compare_element_inner_text(self, locator_definition, expected_text):
        return self.get_element_inner_text(locator_definition) == expected_text

    def create_test_doc(self):
        self.reset_step_counter()
        return docx.Document()

    def add_step_to_test_doc(self, doc_obj, step_definition, associated_screenshot=None, last_step=False):
        self.TEST_STEP_COUNTER = self.TEST_STEP_COUNTER + 1
        doc_obj.add_heading(f"Step {self.TEST_STEP_COUNTER:02d}: {step_definition}")
        if associated_screenshot is not None:
            doc_obj.add_picture(f"{self.DEFAULT_SCREEN_SHOT_PATH}/{associated_screenshot}", width=Inches(6))
        if not last_step and self.TEST_STEP_COUNTER % 2 == 0:
            doc_obj.add_page_break()

    def save_test_doc(self, doc_obj, doc_name):
        doc_obj.save(f"{self.TEST_DOCUMENTS_PATH}/{doc_name}_{get_time_stamp()}.docx")

    @staticmethod
    def compare_lists(list_x, list_y):
        return sorted(list_x) == sorted(list_y)

    def safe_to_proceed(self, key_locator_definition):
        return self.get_number_of_elements(key_locator_definition) > 0

    def start_test_doc(self):
        self.reset_step_counter()
        self.DOC_OBJECT = docx.Document()

    def add_cover_page(self, summary, env_details, tested_by):
        style = self.DOC_OBJECT.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(18)

        p = self.DOC_OBJECT.add_paragraph()
        p.style = self.DOC_OBJECT.styles['Normal']
        p.add_run("SUMMARY: ").bold = True
        p.add_run(summary).italic = True
        p = self.DOC_OBJECT.add_paragraph()
        p.add_run("Environment: ").bold = True
        p.add_run(env_details).italic = True
        p = self.DOC_OBJECT.add_paragraph()
        p.add_run("Tested By: ").bold = True
        p.add_run(tested_by).italic = True
        self.DOC_OBJECT.add_page_break()

    def document_assert_results(self, boolean_result, assert_message, last_assert=False):
        self.update_flow_control_flag(boolean_result)
        self.TEST_STEP_COUNTER = self.TEST_STEP_COUNTER + 1
        self.DOC_OBJECT.add_heading(f"Step {self.TEST_STEP_COUNTER:02d} - {assert_message}: "
                                    f"{boolean_mapper(boolean_result, 'PASSED', 'FAILED')}")
        if not self.SS_FAIL_MODE:
            self.DOC_OBJECT.add_picture(f"{self.DEFAULT_SCREEN_SHOT_PATH}/"
                                        f"{self.get_basic_screenshot()}", width=Inches(6))
        if self.SS_FAIL_MODE and last_assert and self.TEST_STEP_COUNTER % 2 == 0:
            self.DOC_OBJECT.add_page_break()

    def save_doc_results(self, doc_name):
        self.DOC_OBJECT.save(f"{self.TEST_DOCUMENTS_PATH}/{doc_name}_{get_time_stamp()}.docx")
